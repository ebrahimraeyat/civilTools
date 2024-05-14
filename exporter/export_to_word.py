try:
    from docx.shared import Inches
except ImportError:
    package = 'python-docx'
    from freecad_funcs import install_package
    install_package(package_name=package)

from docx import Document
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from docx.enum.text import WD_TEXT_DIRECTION
import os

cfactor_path = os.path.abspath(os.path.join(os.path.dirname(__file__), os.path.pardir))


def get_data_from_model(building, etabs=None):
    if not building:
        return None

    X = building.x_system
    Y = building.y_system

    prop = {'': '',
            'محل اجرای پروژه': building.city,
            'کاربری ساختمان': 'مسکونی',
            'ضریب اهمیت': building.importance_factor,
            'تعداد طبقات': building.number_of_story,
            'ارتفاع ساختمان  )متر(': building.height,
            'سطح خطر نسبی': building.risk_level,
            'شتاب مبنای طرح': building.acc,
            'نوع خاک': building.soil_type}
    if X == Y:
        prop['سیستم سازه ای در دو راستا'] = X.lateral_type
    else:
        prop['سیستم سازه ای در راستای x'] = X.lateral_type
        prop['سیستم سازه ای در راستای y'] = Y.lateral_type

    struc = {'': ('راستای x', 'راستای y'),
             'سیستم سازه': (X.lateral_type, Y.lateral_type),
             'ضریب رفتار': (X.Ru, Y.Ru),
             'ضریب اضافه مقاومت': (X.phi0, Y.phi0),
             'ضریب بزرگنمایی جابجایی': (X.cd, Y.cd),
             'ارتفاع مجاز  )متر(': (X.max_height, Y.max_height)}

    result = {'زمان تناوب تجربی': (building.tx_exp, building.ty_exp),
              'زمان تناوب تحلیلی': (building.tx_an, building.ty_an),
              'ضریب بازتاب': (building.bx, building.by),
              'ضریب زلزله': (building.results[1], building.results[2]),
              'ضریب توزیع در ارتفاع': (building.kx, building.ky),
              'ضریب زلزله دریفت': (building.results_drift[1], building.results_drift[2]),
              'ضریب توزیع در ارتفاع دریفت': (building.kx_drift, building.ky_drift),
              }
    return prop, struc, result

def add_table_figure(
        doc,
        caption: str,
        type_: str='Table ', # 'Figure '
        ):
    paragraph = doc.add_paragraph(type_, style='Caption')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = run = paragraph.add_run()
    r = run._r
    feild_charachter = OxmlElement('w:fldChar')
    feild_charachter.set(qn('w:fldCharType'), 'begin')
    r.append(feild_charachter)
    instr_text = OxmlElement('w:instrText')
    instr_text.text = f' SEQ {type_}\\* ARABIC'
    r.append(instr_text)
    feild_charachter = OxmlElement('w:fldChar')
    feild_charachter.set(qn('w:fldCharType'), 'end')
    r.append(feild_charachter)
    paragraph.add_run(f': {caption} ')

def export(building=None, filename=None, etabs=None, doc=None):
    prop, struc, result = get_data_from_model(building=building, etabs=etabs)
    if doc is None:
        doc = Document(os.path.join(cfactor_path, 'exporter', 'template.docx'))
    table_style = 'List Table 4 Accent 5'
    h = doc.add_heading('محاسبه ضریب زلزله', level=0)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h = doc.add_heading('مشخصات پروژه', level=1)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table_prop = doc.add_table(rows=0, cols=2, style=doc.styles[table_style])
    table_prop.direction = WD_TABLE_DIRECTION.RTL

    for row in range(len(table_prop.rows)):
        for col in range(len(table_prop.columns)):
            cell = table_prop.cell(row, col)
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for key, value in prop.items():
        row_cells = table_prop.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    h = doc.add_heading('مشخصات سیستم سازه ای', level=1)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    struc_table = doc.add_table(rows=0, cols=3, style=doc.styles[table_style])
    for key, value in struc.items():
        row_cells = struc_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value[0])
        row_cells[2].text = str(value[1])

    h = doc.add_heading('ضریب زلزله', level=1)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    result_table = doc.add_table(rows=1, cols=3, style=doc.styles[table_style])
    hdr_cells = result_table.rows[0].cells
    hdr_cells[1].text = 'راستای x'
    hdr_cells[2].text = 'راستای y'
    for key, value in result.items():
        row_cells = result_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = f'{value[0]:.3f}'
        row_cells[2].text = f'{value[1]:.3f}'

    if filename:
        doc.save(filename)
    return doc
