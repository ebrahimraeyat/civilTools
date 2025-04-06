from pathlib import Path
import random
import colorsys
import json
from typing import Union

import pandas as pd
from PySide2.QtCore import QAbstractTableModel, Qt, QModelIndex, QItemSelection
from PySide2.QtGui import QColor, QIcon
from PySide2 import QtCore, QtWidgets

import civiltools_rc
# import matplotlib.cm as cm
# from matplotlib.colors import Normalize
# import matplotlib
import FreeCAD

from freecad_funcs import open_file, get_color

civiltools_path = Path(__file__).absolute().parent

param = FreeCAD.ParamGet("User parameter:BaseApp/Preferences/Mod/civilTools")

low = get_color(param, 'low_result_colors', 16777215)
intermediate = get_color(param, 'medium_result_colors', 4294934527)
high = get_color(param, 'high_result_colors', 4283793407)
open_results_file = param.GetBool('open_results_file', True)
# def color_map_color(value, norm, cmap_name='rainbow'):
#     cmap = cm.get_cmap(cmap_name)  # PiYG
#     rgb = cmap(norm(abs(value)))[:3]  # will return rgba, we take only first 3 so we get rgb
#     color = matplotlib.colors.rgb2hex(rgb)
#     return color

class PandasModel(QAbstractTableModel):
    '''
    MetaClass Model for showing Results from pandas dataframe
    '''
    check_states_bool = {False: 0, True: 2}
    check_states_numeric = {0: False, 2: True}

    def __init__(self,
                 data,
                 kwargs: Union[dict, None] = None,
                 negative_value: bool=False,
                 ):
        QAbstractTableModel.__init__(self)
        self.df = data
        self.kwargs = kwargs
        self.negative_value = negative_value

    def rowCount(self, parent=None):
        return self.df.shape[0]

    def columnCount(self, parent=None):
        return self.df.shape[1]

    def headerData(self, col, orientation, role):
        if role != Qt.DisplayRole:
            return
        if orientation == Qt.Horizontal:
            return self.df.columns[col]
        return int(col + 1)

    def data(self, index, role=Qt.DisplayRole):
        if not index.isValid():
            return
        col = index.column()
        row = index.row()
        value = self.df.iat[row, col]
        if role == Qt.CheckStateRole and self.df.dtypes[col] == bool:
            return self.check_states_bool.get(bool(value), 1)
        elif role == Qt.DisplayRole and self.df.dtypes[col] != bool:
            return str(value)
        
    def setData(self, index, value, role=Qt.EditRole):
        if not index.isValid():
            return False
        col = index.column()
        row = index.row()
        if role == Qt.CheckStateRole and self.df.dtypes[col] == bool:
            self.df.iloc[row, col] = self.check_states_numeric.get(int(value))
            self.dataChanged.emit(index, index)
            return True
        elif role == Qt.EditRole:
            if pd.api.types.is_numeric_dtype(self.df.iloc[:, col]):
                if not self.negative_value and float(value) < 0:
                    return False
                else:
                    self.df.iloc[row, col] = float(value)
                self.dataChanged.emit(index, index)
                return True
            elif pd.api.types.is_string_dtype(self.df.iloc[:, col]):
                self.df.iloc[row, col] = str(value)
                self.dataChanged.emit(index, index)
                return True
        return False


class DriftModel(PandasModel):
    def __init__(self, df, kwargs=None):
        super(DriftModel, self).__init__(df, kwargs)
        max_drift = 'Max Drift'
        avg_drift = 'Avg Drift'
        allowable_drift = 'Allowable Drift'
        self.df = self.df[[
            'Story',
            'OutputCase',
            max_drift,
            avg_drift,
            allowable_drift,
        ]]
        self.df = self.df.astype({max_drift: float, avg_drift: float, allowable_drift: float})
        headers = tuple(self.df.columns)
        self.max_i = headers.index(max_drift)
        self.avg_i = headers.index(avg_drift)
        self.allow_i = headers.index(allowable_drift)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            allow_drift = float(self.df.iloc[row][self.allow_i])
            if role == Qt.DisplayRole:
                if col in (self.max_i, self.avg_i, self.allow_i):
                    return f"{value:.4f}"
                return value
            elif role == Qt.BackgroundColorRole:
                if col in (self.avg_i, self.max_i):
                    if float(value) > allow_drift:
                        return QColor(*high)
                    else:
                        return QColor(*low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None

class TorsionModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(TorsionModel, self).__init__(data, kwargs)
        headers = [
            'Story',
            'Label',
            'OutputCase',
            'Max Drift',
            'Avg Drift',
            'Ratio',
        ]
        self.df = self.df[headers]
        i_story = headers.index('Story')
        i_label = headers.index('Label')
        self.i_ratio = headers.index('Ratio')
        self.i_max = headers.index('Max Drift')
        self.i_avg = headers.index('Avg Drift')
        self.col_function = (i_story, i_label)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_max, self.i_avg, self.i_ratio):
                    return f"{value:.4f}"
                return str(value)
            elif role == Qt.BackgroundColorRole:
                value = float(self.df.iloc[row][self.i_ratio])
                # if col == i_ratio:
                    # value = float(value)
                if value <= 1.2:
                    return QColor(*low)
                elif 1.2 < value < 1.4:
                    return QColor(*intermediate)
                elif value > 1.4:
                    return QColor(*high)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None

class BaseShearModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(BaseShearModel, self).__init__(data, kwargs)
        self.i_ratio = list(self.df.columns).index('Ratio')
        self.i_scale = list(self.df.columns).index('Scale')

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col == 0:
                    return str(value)
                elif col in (self.i_ratio, self.i_scale):
                    return f"{value:.2f}"
                else:
                    return f"{value:.0f}"
                
            elif role == Qt.BackgroundColorRole:
                return QColor(*low)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None


class StoryForcesModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(StoryForcesModel, self).__init__(data, kwargs)
        self.df = self.df[[
            'Story',
            'OutputCase',
            'VX',
            'VY',
            'Vx %',
            'Vy %',
        ]]
        headers = tuple(self.df.columns)
        self.i_vx = headers.index('Vx %')
        self.i_vy = headers.index('Vy %')

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                return str(value)
            elif role == Qt.BackgroundColorRole:
                fx_percentage = float(self.df.iloc[row][self.i_vx])
                fy_percentage = float(self.df.iloc[row][self.i_vy])
                if max(fx_percentage, fy_percentage) >= .35:
                    return QColor(*intermediate)
                else:
                    return QColor(*low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None


class ColumnsRatioModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(ColumnsRatioModel, self).__init__(data, kwargs)
        all_cols = list(self.df)
        self.df[all_cols] = self.df[all_cols].astype(str)
        headers = tuple(self.df.columns)
        i_story = headers.index('Story')
        i_label = headers.index('Label')
        self.col_function = (i_label, i_story)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                return str(value)
            elif role == Qt.BackgroundColorRole:
                ratio = float(self.df.iloc[row]['Ratio'])
                if ratio > 1:
                    return QColor(*high)
                else:
                    return QColor(*low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None


class BeamsRebarsModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(BeamsRebarsModel, self).__init__(data, kwargs)
        all_cols = list(self.df)
        self.df[all_cols] = self.df[all_cols].astype(str)
        headers = tuple(self.df.columns)
        i_story = headers.index('Story')
        i_label = headers.index('Label')
        self.i_location = headers.index('location')
        self.i_ta1 = headers.index('Top Area1')
        self.i_ta2 = headers.index('Top Area2')
        self.i_ba1 = headers.index('Bot Area1')
        self.i_ba2 = headers.index('Bot Area2')
        self.i_v1 = headers.index('VRebar1')
        self.i_v2 = headers.index('VRebar2')
        self.col_function = (i_label, i_story)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (
                            self.i_ta1,
                            self.i_ta2,
                            self.i_ba1,
                            self.i_ba2,
                            ):
                    value = round(float(value), 1)
                if col in (
                            self.i_v1,
                            self.i_v2,
                            ):
                    value = round(float(value) * 100, 1)
                if col == self.i_location:
                    value = int(float(value))
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_ta1, self.i_ta2):
                    if float(self.df.iloc[row][self.i_ta2]) > float(self.df.iloc[row][self.i_ta1]) * 1.02:
                        return QColor(*high)
                    else:
                        return QColor(*low)
                if col in (self.i_ba1, self.i_ba2):
                    if float(self.df.iloc[row][self.i_ba2]) > float(self.df.iloc[row][self.i_ba1]) * 1.02:
                        return QColor(*high)
                    else:
                        return QColor(*low)
                if col in (self.i_v1, self.i_v2):
                    if float(self.df.iloc[row][self.i_v2]) > float(self.df.iloc[row][self.i_v1]) * 1.02:
                        return QColor(*high)
                    else:
                        return QColor(*low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None


class IrregularityOfMassModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(IrregularityOfMassModel, self).__init__(data, kwargs)
        headers = list(self.df.columns)
        self.df[headers] = self.df[headers].astype(str)
        self.i_mass_x = headers.index('Mass X')
        self.i_below = headers.index('1.5 * Below')
        self.i_above = headers.index('1.5 * Above')
        self.i_story = headers.index('Story')
        # self.col_function = (0, 4)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col == self.i_story:
                    return str(value)
                return f'{float(value):.0f}'
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_below, self.i_above):
                    if float(self.df.iloc[row][self.i_mass_x]) > \
                        float(self.df.iloc[row][col]):
                        return QColor(*high)
                    else:
                        return QColor(*low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None

class StoryStiffnessModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(StoryStiffnessModel, self).__init__(data, kwargs)
        headers = list(self.df.columns)
        self.df[headers] = self.df[headers].astype(str)
        self.i_kx = headers.index('Kx')
        self.i_ky = headers.index('Ky')
        self.i_kx_above = headers.index('Kx / kx+1')
        self.i_ky_above = headers.index('Ky / ky+1')
        self.i_kx_3above = headers.index('Kx / kx_3ave')
        self.i_ky_3above = headers.index('Ky / ky_3ave')
        # self.col_function = (0, 4)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (
                    self.i_kx_above,
                    self.i_kx_3above,
                    self.i_ky_above,
                    self.i_ky_3above,
                    ):
                    if value == '-':
                        return value
                    return f'{float(value):.3f}'
                elif col in (
                    self.i_kx,
                    self.i_ky,
                    ):
                    return f'{float(value):.0f}'
                return value
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_kx_above, self.i_ky_above):
                    k = self.df.iloc[row][col]
                    return self.get_color(k, .6, .7)
                elif col in (self.i_kx_3above, self.i_ky_3above):
                    k = self.df.iloc[row][col]
                    return self.get_color(k, .7, .8)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None

    @staticmethod
    def get_color(k, a, b):
        if k == '-':
            return None
        k = float(k)
        if k < a:
            return QColor(*high)
        elif k < b:
            return QColor(*intermediate)
        else:
            return QColor(*low)


class BeamsJModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(BeamsJModel, self).__init__(data, kwargs)
        headers = tuple(self.df.columns)
        self.i_T = headers.index('T')
        self.i_Tcr = headers.index('phi_Tcr')
        self.i_j = headers.index('j')
        self.i_init_j = headers.index('init_j')
        self.col_function = (2,)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_T, self.i_j, self.i_Tcr, self.i_init_j):
                    return f'{value:.2f}'
                return str(value)
            elif role == Qt.BackgroundColorRole:
                j = self.df.iloc[row][self.i_j]
                j_init = self.df.iloc[row][self.i_init_j]
                # if col == i_ratio:
                    # value = float(value)
                if j == j_init:
                    return QColor(*low)
                else:
                    return QColor(*intermediate)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return False

class HighPressureColumnModel(PandasModel):
    def __init__(self, data, kwargs=None):
        super(HighPressureColumnModel, self).__init__(data, kwargs)
        headers = tuple(self.df.columns)
        self.i_p = headers.index('P')
        self.i_t2 = headers.index('t2')
        self.i_t3 = headers.index('t3')
        self.i_fc = headers.index('fc')
        self.i_Agfc = headers.index('limit*Ag*fc')
        self.i_hp = headers.index('Result')
        self.col_function = (3,)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_p, self.i_t2, self.i_t3, self.i_fc, self.i_Agfc):
                    return f'{value:.0f}'
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if self.df.iloc[row][self.i_hp]:
                    return QColor(*intermediate)
                else:
                    return QColor(*low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None

class Column100_30Model(PandasModel):
    def __init__(self, data, kwargs=None):
        super(Column100_30Model, self).__init__(data, kwargs)
        headers = tuple(self.df.columns)
        self.i_p = headers.index('P')
        self.i_mmajor = headers.index('MMajor')
        self.i_mminor = headers.index('MMinor')
        self.i_ratio = headers.index('Ratio')
        self.i_result = headers.index('Result')
        self.col_function = (2,)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_p, self.i_ratio):
                    return f'{value:.2f}'
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if self.df.iloc[row][self.i_result]:
                    return QColor(*low)
                else:
                    return QColor(*high)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None


class JointShearBCC(PandasModel):
    def __init__(self, data, kwargs=None):
        super(JointShearBCC, self).__init__(data, kwargs)
        headers = tuple(self.df.columns)
        try:
            self.i_maj_js = headers.index('JSMajRatio')
            self.i_min_js = headers.index('JSMinRatio')
        except ValueError:
            self.i_maj_js = -2
            self.i_min_js = -2
        try:
            self.i_maj_bc = headers.index('BCMajRatio')
            self.i_min_bc = headers.index('BCMinRatio')
        except ValueError:
            self.i_maj_bc = -2
            self.i_min_bc = -2
        self.col_function = (2,)

    def data(self, index, role=Qt.DisplayRole):
        if index.isValid():
            row = index.row()
            col = index.column()
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_maj_js, self.i_min_js, self.i_maj_bc, self.i_min_bc):
                    try:
                        value = float(value)
                        return f'{value:.2f}'
                    except (ValueError, TypeError):
                        return str(value)
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_maj_js, self.i_min_js, self.i_maj_bc, self.i_min_bc):
                    try:
                        value = float(value)
                        if value <= 1:
                            return QColor(*low)
                        else:
                            return QColor(*high)
                    except (ValueError, TypeError):
                        return QColor(*intermediate)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None


class ExpandLoadSets(PandasModel):
    def __init__(self, data, kwargs=None):
        super(ExpandLoadSets, self).__init__(data, kwargs)
        headers = tuple(self.df.columns)
        self.i_uniquename = headers.index('UniqueName')
        unique_names = self.df['UniqueName'].unique()
        self.colors = {}
        for name in unique_names:
            self.colors[name] = random_color()

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                return f'{value}'
            elif role == Qt.BackgroundColorRole:
                name = self.df.iloc[row][self.i_uniquename]
                return QColor.fromRgb(*self.colors[name])
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            
    def setData(self, index, value, role=Qt.EditRole):
        return None
    

class BeamDeflectionTableModel(PandasModel):

    def __init__(self, df, kwargs=None, parent=None):
        '''
        beam_data : dict with keys = beam_name and value is dict of properties
        '''
        super().__init__(df, kwargs)
        self.col_function = (0,)
        self.etabs = None
        self.results = None
        self.console_short_term = None
        self.console_long_term = None
        self.continues_short_term = None
        self.continues_long_term = None
        
    def data(self, index, role=Qt.DisplayRole):
        if not index.isValid():
            return
        col = index.column()
        row = index.row()
        col_name = self.df.columns[col]
        value = self.df.iat[row, col]
        if role == Qt.CheckStateRole and self.df.dtypes[col] == bool:
            return self.check_states_bool.get(bool(value), 1)
        elif role == Qt.DisplayRole and self.df.dtypes[col] != bool:
            return str(self.df.iat[index.row(), index.column()])
        elif role == Qt.BackgroundColorRole:
            if self.results is not None:
                beam_name = str(self.df.iat[row, 0])
                is_console = self.df['Console'].iloc[row]
                if is_console:
                    short_term = self.console_short_term
                    long_term = self.console_long_term
                else:
                    short_term = self.continues_short_term
                    long_term = self.continues_long_term
                def1 = self.results[0][row]
                def2 = self.results[1][row]
                minus_length = self.df['Minus Length'].iloc[row]
                ln = self.etabs.frame_obj.get_length_of_frame(beam_name) - minus_length
                short_term, long_term = self.get_deflection_check_result(def1, def2, ln, short_term, long_term)
                if short_term and long_term:
                    return QColor(*low)
                elif not short_term:
                    return QColor(*high)
                elif not long_term:
                    return QColor(*intermediate)
            elif col_name in ('Width', 'Height') and self.df.iloc[row, col] <= 0:
                return QColor('yellow')
            
    def get_deflection_check_result(self,
        def1: float,
        def2: float,
        ln: float,
        short_term: float=360,
        long_term: float=480,
        ):
        allow_def1 = ln / short_term
        allow_def2 = ln / long_term
        if def1 <= allow_def1:
            short_term = True
        else:
            short_term = False
        # combo 2
        if def2 <= allow_def2:
            long_term = True
        else:
            long_term = False
        return short_term, long_term

    def flags(self, index):
        if not index.isValid():
            return Qt.ItemIsEnabled
        col = index.column()
        if self.df.dtypes[col] == bool:
            return Qt.ItemIsEnabled | Qt.ItemIsUserCheckable
        elif pd.api.types.is_numeric_dtype(self.df.iloc[:, col]):
            return Qt.ItemFlags(
                QAbstractTableModel.flags(self, index) | Qt.ItemIsEditable)
        elif pd.api.types.is_string_dtype(self.df.iloc[:, col]):
            return Qt.ItemIsSelectable | Qt.ItemIsEnabled
        
            # if role == Qt.DecorationRole:
        #     value = self._data[index.row()][index.column()]
        #     if isinstance(value, float):
        #         return QtGui.QIcon('calendar.png')
        # return None

    # def sort(self, col, order):
    #     """Sort table by given column number."""
    #     self.layoutAboutToBeChanged.emit()
    #     self.df.sort_values(
    #         by=self.df.columns[col],
    #         ascending = order == Qt.AscendingOrder,
    #         kind="mergesort",
    #         inplace=True,
    #     )
    #     self.df.reset_index(drop=True, inplace=True)
    #     self.layoutChanged.emit()
            

class ResultWidget(QtWidgets.QDialog):
    # main widget for user interface
    def __init__(self,
                 data,
                 model,
                 function=None,
                 parent=None,
                 kwargs: Union[None, dict]=None,
                ):
        super(ResultWidget, self).__init__(parent)
        self.setObjectName('result_widget')
        self.push_button_to_excel = QtWidgets.QPushButton()
        self.push_button_to_excel.setIcon(QIcon(str(civiltools_path / 'images' / 'xlsx.png')))
        self.push_button_to_word = QtWidgets.QPushButton()
        self.push_button_to_word.setIcon(QIcon(str(civiltools_path / 'images' / 'word.png')))
        self.open_exported_file = QtWidgets.QCheckBox("Open")
        self.open_exported_file.setChecked(open_results_file)
        label = QtWidgets.QLabel("Filter")
        self.lineEdit = QtWidgets.QLineEdit()
        label2 = QtWidgets.QLabel("By Column:")
        self.comboBox = QtWidgets.QComboBox()
        hbox = QtWidgets.QHBoxLayout()
        hbox.addWidget(label)
        hbox.addWidget(self.lineEdit)
        hbox.addWidget(label2)
        hbox.addWidget(self.comboBox)
        hbox.addWidget(self.push_button_to_word)
        hbox.addWidget(self.push_button_to_excel)
        hbox.addWidget(self.open_exported_file)
        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addLayout(hbox)
        self.result_table_view = QtWidgets.QTableView()
        self.result_table_view.setSortingEnabled(True)
        self.table_layout = QtWidgets.QHBoxLayout()
        self.table_layout.addWidget(self.result_table_view)
        self.vbox.addLayout(self.table_layout)
        self.setLayout(self.vbox)
        self.function = function
        self.data = data
        self.model = model(data, kwargs)
        # self.result_table_view.setModel(self.model)
        self.proxy = QtCore.QSortFilterProxyModel(self)
        self.proxy.setSourceModel(self.model)
        self.result_table_view.setModel(self.proxy)
        if kwargs is not None:
            custom_delegate = kwargs.get('custom_delegate', None)
            if custom_delegate is not None:
                self.result_table_view.setItemDelegate(custom_delegate(self))
        self.comboBox.addItems(self.model.df.columns)
        self.lineEdit.textChanged.connect(self.on_lineEdit_textChanged)
        self.comboBox.currentIndexChanged.connect(self.on_comboBox_currentIndexChanged)
        # self.horizontalHeader = self.result_table_view.horizontalHeader()
        # self.horizontalHeader.sectionClicked.connect(self.on_view_horizontalHeader_sectionClicked)
        self.push_button_to_excel.clicked.connect(self.export_to_excel)
        self.push_button_to_word.clicked.connect(self.export_to_word)
        self.resize_columns()
        if self.function:
            self.result_table_view.clicked.connect(self.row_clicked)
        

    def get_current_row_col(self, index=None):
        '''
        get an index and return the actual source row and col 
        '''
        if index is None:
            index = self.result_table_view.currentIndex()
        source_index = self.proxy.mapToSource(index)
        return source_index.row(), source_index.column()

    def row_clicked(self, index):
        row, _ = self.get_current_row_col(index)
        args = []
        for col in self.model.col_function:
            value = str(self.model.data(self.model.index(row, col)))
            args.append(value)
        self.function(*args)

    def save_table_to_json(self,
                           filename: str='',
                           ):
        
        data = []
        # Header data
        for col in range(self.model.columnCount()):
            text = self.model.headerData(col, Qt.Horizontal, Qt.DisplayRole)
            data.append({"row": 0, "col": col, "text": text, "color": ""})
        for row in range(self.model.rowCount()):
            for col in range(self.model.columnCount()):
                index = self.model.index(row, col)
                text = self.model.data(index)
                color = self.model.data(index, Qt.BackgroundColorRole)
                if color is None:
                    color = ""
                else:
                    color = color.name()
                data.append({"row": row + 1, "col": col, "text": text, "color": color})
        with open(filename, "w") as f:
            json.dump(data, f, indent=4)

    def export_to_excel(self):
        filename, _ = QtWidgets.QFileDialog.getSaveFileName(self, 'export to excel',
                                                  '', "excel(*.xlsx)")
        if filename == '':
            return
        try:
            import jinja2
        except ModuleNotFoundError:
            import subprocess
            package = 'Jinja2'
            subprocess.check_call(['python', "-m", "pip", "install", package])
            css_alt_rows = 'background-color: powderblue; color: black;'
            css_indexes = 'background-color: steelblue; color: white;'
            import numpy as np
            (self.model.df.style.apply(lambda col: np.where(col.index % 2, css_alt_rows, None)) # alternating rows
                    .applymap_index(lambda _: css_indexes, axis=0) # row indexes (pandas 1.4.0+)
                    .applymap_index(lambda _: css_indexes, axis=1) # col indexes (pandas 1.4.0+)
            ).to_excel(filename, engine='openpyxl')
        else:
            with pd.ExcelWriter(filename) as writer:
                    self.model.df.to_excel(writer)
        if self.open_exported_file.isChecked():
            open_file(filename)

    def export_to_word(self,
                       ali='',
                       doc=None,
                       ):
        filename, _ = QtWidgets.QFileDialog.getSaveFileName(self, 'export to word',
                                                  '', "word(*.docx)")
        if filename == '':
            return
        try:
            from docx.shared import Inches
        except ImportError:
            package = 'python-docx'
            from freecad_funcs import install_package
            install_package(package)
    
        from docx import Document
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        # create a new document
        if doc is None:
            doc = Document()
        table_docx = doc.add_table(rows=self.model.rowCount()+1, cols=self.model.columnCount())
        table_docx.style = 'Table Grid'

        # write the column headers to the first row of the table
        for j in range(self.model.columnCount()):
            cell = table_docx.cell(0, j)
            cell.text = self.model.headerData(j, Qt.Horizontal, Qt.DisplayRole)
            # Set header text to bold
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.italic = True
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "#244061"))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # write the data to the remaining rows of the table
        for row in range(self.model.rowCount()):
            for col in range(self.model.columnCount()):
                index = self.model.index(row, col)
                text = self.model.data(index)
                color = self.model.data(index, Qt.BackgroundColorRole)
                cell = table_docx.cell(row + 1, col)
                cell.text = text
                if color:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color.name()))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
        # save the document
        doc.save(filename)
        if self.open_exported_file.isChecked():
            open_file(filename)

    def resize_columns(self):
        self.result_table_view.resizeColumnsToContents()
        self.adjustSize()
        # width = 0
        # for col in range(len(self.model.df.columns)):
        #     width += self.result_table_view.columnWidth(col)
        # self.result_table_view.setFixedWidth(width)

    # @QtCore.Slot(int)
    # def on_view_horizontalHeader_sectionClicked(self, logicalIndex):
    #     self.logicalIndex   = logicalIndex
    #     self.menuValues     = QtWidgets.QMenu(self)
    #     self.signalMapper   = QtCore.QSignalMapper(self)  

    #     self.comboBox.blockSignals(True)
    #     self.comboBox.setCurrentIndex(self.logicalIndex)
    #     self.comboBox.blockSignals(True)


    #     valuesUnique = list(self.model.df.iloc[:, self.logicalIndex].unique())

    #     actionAll = QtWidgets.QAction("All", self)
    #     actionAll.triggered.connect(self.on_actionAll_triggered)
    #     self.menuValues.addAction(actionAll)
    #     self.menuValues.addSeparator()

    #     for actionNumber, actionName in enumerate(sorted(list(set(valuesUnique)))):              
    #         action = QtWidgets.QAction(actionName, self)
    #         self.signalMapper.setMapping(action, actionNumber)  
    #         action.triggered.connect(self.signalMapper.map)  
    #         self.menuValues.addAction(action)

    #     self.signalMapper.mapped.connect(self.on_signalMapper_mapped)  

    #     headerPos = self.result_table_view.mapToGlobal(self.horizontalHeader.pos())        

    #     posY = headerPos.y() + self.horizontalHeader.height()
    #     posX = headerPos.x() + self.horizontalHeader.sectionPosition(self.logicalIndex)

    #     self.menuValues.exec_(QtCore.QPoint(posX, posY))

    @QtCore.Slot()
    def on_actionAll_triggered(self):
        filterColumn = self.logicalIndex
        filterString = QtCore.QRegExp(  "",
                                        QtCore.Qt.CaseInsensitive,
                                        QtCore.QRegExp.RegExp
                                        )

        self.proxy.setFilterRegExp(filterString)
        self.proxy.setFilterKeyColumn(filterColumn)

    @QtCore.Slot(int)
    def on_signalMapper_mapped(self, i):
        stringAction = self.signalMapper.mapping(i).text()
        filterColumn = self.logicalIndex
        filterString = QtCore.QRegExp(  stringAction,
                                        QtCore.Qt.CaseSensitive,
                                        QtCore.QRegExp.FixedString
                                        )

        self.proxy.setFilterRegExp(filterString)
        self.proxy.setFilterKeyColumn(filterColumn)

    @QtCore.Slot(str)
    def on_lineEdit_textChanged(self, text):
        search = QtCore.QRegExp(    text,
                                    QtCore.Qt.CaseInsensitive,
                                    QtCore.QRegExp.RegExp
                                    )

        self.proxy.setFilterRegExp(search)

    @QtCore.Slot(int)
    def on_comboBox_currentIndexChanged(self, index):
        self.proxy.setFilterKeyColumn(index)
        

class ControlColumnResultWidget(ResultWidget):

    def __init__(self,
                data,
                model,
                function=None,
                parent=None,
                kwargs: Union[None, dict]=None,
            ):
        super(ControlColumnResultWidget, self).__init__(data, model, function, parent, kwargs)
        self.result_table_view.setSortingEnabled(False)
        self.add_legend()
        # self.result_table_view.setItemDelegate(ColumnsControlDelegate(self))
    
    def add_legend(self):
        from qt_models.columns_control_models import CompareTwoColumnsColorEnum
        from prop_frame import CompareTwoColumnsEnum
        vgrid = QtWidgets.QVBoxLayout()
        for color_enum in CompareTwoColumnsColorEnum:
            text = f"{CompareTwoColumnsEnum(color_enum.value).name}"
            texts = text.split("_")
            text = " ".join([t.capitalize() for t in texts])
            label = QtWidgets.QLabel(text)
            label.setStyleSheet(f"background-color: {color_enum.name};")
            label.setAlignment(Qt.AlignCenter)
            vgrid.insertWidget(color_enum.value, label)
        self.table_layout.addLayout(vgrid)


    def row_clicked(self, index):
        row, col = self.get_current_row_col(index)
        col_label = self.model.headerData(col, orientation=Qt.Horizontal, role=Qt.DisplayRole)
        story = self.model.headerData(row, orientation=Qt.Vertical, role=Qt.DisplayRole)
        args = [col_label, story]
        self.function(*args)

class ExpandedLoadSetsResults(ResultWidget):
    def __init__(self, data, model, function=None, parent=None):
        super(ExpandedLoadSetsResults, self).__init__(data, model, function, parent)
        self.cancel_pushbutton = QtWidgets.QPushButton()
        self.cancel_pushbutton.setIcon(QIcon(str(civiltools_path / 'images' / 'cancel.svg')))
        self.cancel_pushbutton.setText('&Cancel')
        self.apply_pushbutton = QtWidgets.QPushButton()
        self.apply_pushbutton.setIcon(QIcon(str(civiltools_path / 'images' / 'etabs.png')))
        self.apply_pushbutton.setText('&Apply')

        hbox = QtWidgets.QHBoxLayout()
        hbox.addWidget(self.cancel_pushbutton)
        hbox.addWidget(self.apply_pushbutton)
        self.vbox.addLayout(hbox)


def show_results(
        data,
        model,
        function=None,
        etabs=None,
        json_file_name:str='',
        kwargs: Union[None, dict]=None,
        result_widget=ResultWidget,
        ):
    win = result_widget(data, model, function, kwargs=kwargs)
    # Gui.Control.showDialog(win)
    mdi = get_mdiarea()
    if not mdi:
        return None
    sub = mdi.addSubWindow(win)
    if json_file_name:
        suffix = '.json'
        if json_file_name.endswith(suffix):
            json_len = len(suffix)
            title = json_file_name[:-json_len]
        else:
            title = json_file_name
        sub.setWindowTitle(title)
    sub.show()
    # Save table as json
    if etabs is not None:
        filename = etabs.get_json_file_path_for_table_results(json_filename=json_file_name)
    else:
        filename = ""
    if filename and isinstance(filename, Path):
        if not filename.parent.exists():
            filename.parent.mkdir(parents=True)
        win.save_table_to_json(filename=filename)


def get_mdiarea():
    """ Return FreeCAD MdiArea. """
    import FreeCADGui as Gui
    import PySide2
    mw = Gui.getMainWindow()
    if not mw:
        return None
    childs = mw.children()
    for c in childs:
        if isinstance(c, PySide2.QtWidgets.QMdiArea):
            return c
    return None

def random_color():
    h,s,l = random.random(), 0.5 + random.random()/2.0, 0.4 + random.random()/5.0
    return [int(256*i) for i in colorsys.hls_to_rgb(h,l,s)]


if __name__ == "__main__":
        import sys
        data = {'Name': '1',
        'is_console': False,
        'Label': 'B12',
        'Story': 'STORY1',
        'minus_length': 30,
        'add_torsion_rebar': False,
        'add_rebar': 5,
        'cover': 4,
        'width': 40,
        'height': 50,
        'result': ''}
        app=QtWidgets.QApplication(sys.argv)
        df = pd.DataFrame.from_dict(10*[data])
        new_rows = pd.DataFrame({'Name': ['2', '3'], 'is_console': [False] * 2,
        'Label': ['B12'] * 2,
        'Story': ['STORY'] * 2,
        'minus_length': [30] * 2,
        'add_torsion_rebar':[False] * 2,
        'add_rebar': [5] * 2,
        'cover': [4] * 2,
        'width': [40] * 2,
        'height': [50] * 2,
        'result': [''] * 2,
        })
        df = df.append(new_rows, ignore_index=True)
        widget = ResultWidget(df, model=BeamDeflectionTableModel)
        widget.show()
        app.exec_()